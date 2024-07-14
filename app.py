from flask import Flask, request, render_template, send_file
import zipfile
import os
import shutil
import docx
from docx import Document
import graphviz

app = Flask(__name__)

def extract_vba_code(file):
    with zipfile.ZipFile(file, 'r') as zip_ref:
        zip_ref.extractall('temp')
    vba_code = ""
    for root, dirs, files in os.walk('temp'):
        for file in files:
            if file.endswith('.vb') or file.endswith('.bas') or file.endswith('.cls'):
                with open(os.path.join(root, file), 'r') as vb_file:
                    vba_code += vb_file.read()
    shutil.rmtree('temp')
    return vba_code

def generate_documentation(vba_code):
    doc = Document()
    doc.add_heading('VBA Macro Documentation', 0)
    doc.add_heading('Extracted VBA Code:', level=1)
    doc.add_paragraph(vba_code)

    filename = "VBA_Documentation.docx"
    doc.save(filename)
    return filename

def generate_flow_diagram(vba_code):
    dot = graphviz.Digraph(comment='VBA Macro Flow Diagram')

    # Placeholder for logic extraction and diagram creation
    # This part needs to be customized based on the actual structure and logic of your VBA code
    functions = extract_functions(vba_code)
    for func in functions:
        dot.node(func, func)
    for func in functions:
        for call in extract_function_calls(vba_code, func):
            dot.edge(func, call)

    diagram_path = 'flow_diagram'
    dot.render(diagram_path, format='png')
    return f'{diagram_path}.png'

def extract_functions(vba_code):
    # Extract function names from the VBA code
    functions = []
    lines = vba_code.splitlines()
    for line in lines:
        if line.strip().lower().startswith("sub "):
            func_name = line.strip().split()[1].split('(')[0]
            functions.append(func_name)
    return functions

def extract_function_calls(vba_code, func):
    # Extract function calls within a function from the VBA code
    calls = []
    lines = vba_code.splitlines()
    inside_func = False
    for line in lines:
        if line.strip().lower().startswith(f"sub {func.lower()}"):
            inside_func = True
        elif line.strip().lower().startswith("end sub"):
            inside_func = False
        elif inside_func and "(" in line and ")" in line:
            call = line.strip().split('(')[0].split()[-1]
            calls.append(call)
    return calls

@app.route('/')
def upload_file():
    return render_template('upload.html')

@app.route('/analyze', methods=['POST'])
def analyze_file():
    file = request.files['file']
    file.save("uploaded_file.xlsm")
    vba_code = extract_vba_code("uploaded_file.xlsm")
    documentation_file = generate_documentation(vba_code)
    diagram_file = generate_flow_diagram(vba_code)
    return render_template('result.html', documentation_file=documentation_file, diagram_file=diagram_file)

if __name__ == '__main__':
    app.run(debug=True)
